import docx
doc= docx.Document('C:\\python\\word.docx')
print(len(doc.paragraphs))

for val in range(len(doc.paragraphs)):
    print(doc.paragraphs[val].text)
    print('--------------------')